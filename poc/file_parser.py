"""
file_parser.py – Extract text from uploaded attachments
=========================================================
Supported types:
  - PDF         → pypdf (already installed)
  - DOCX        → python-docx (already installed)
  - XLSX/XLSM   → openpyxl (already installed)
  - MSG (Outlook) → extract-msg (extracts subject + from + date + FULL BODY)
  - EML (plain email) → Python stdlib email parser
  - Images (PNG/JPG/WEBP/GIF) → Claude vision API
  - TXT/CSV     → plain read
"""

import io
import os
import base64
import tempfile
from pathlib import Path

SUPPORTED_EXTENSIONS = {
    ".pdf", ".docx", ".doc",
    ".xlsx", ".xlsm", ".xls", ".csv",
    ".msg", ".eml",
    ".png", ".jpg", ".jpeg", ".webp", ".gif",
    ".txt",
}

MAX_TEXT_CHARS = 8000   # clip per file before sending to LLM


# ── PDF ───────────────────────────────────────────────────────────────────────

def _parse_pdf(data: bytes) -> str:
    from pypdf import PdfReader
    reader = PdfReader(io.BytesIO(data))
    pages  = []
    for page in reader.pages[:20]:       # cap at 20 pages
        text = page.extract_text() or ""
        if text.strip():
            pages.append(text.strip())
    return "\n\n".join(pages)


# ── DOCX ──────────────────────────────────────────────────────────────────────

def _parse_docx(data: bytes) -> str:
    from docx import Document
    doc  = Document(io.BytesIO(data))
    text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
    # Also grab table cells
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                text += "\n" + " | ".join(cells)
    return text


# ── XLSX / XLSM ───────────────────────────────────────────────────────────────

def _parse_xlsx(data: bytes) -> str:
    import openpyxl
    wb    = openpyxl.load_workbook(io.BytesIO(data), read_only=True, data_only=True)
    lines = []
    for sheet_name in wb.sheetnames[:5]:          # cap at 5 sheets
        ws = wb[sheet_name]
        lines.append(f"[Sheet: {sheet_name}]")
        row_count = 0
        for row in ws.iter_rows(values_only=True):
            cells = [str(c) for c in row if c is not None and str(c).strip()]
            if cells:
                lines.append(" | ".join(cells))
                row_count += 1
            if row_count >= 100:                   # cap rows per sheet
                break
    wb.close()
    return "\n".join(lines)


# ── MSG (Outlook) ─────────────────────────────────────────────────────────────

def _parse_msg(data: bytes) -> str:
    """
    Parse an Outlook .msg file. extract_msg only accepts a path, not raw bytes,
    so we write a temp file first.
    """
    import extract_msg
    import re as _re

    with tempfile.NamedTemporaryFile(suffix=".msg", delete=False) as tf:
        tf.write(data)
        tmp_path = tf.name
    try:
        msg = extract_msg.Message(tmp_path)

        subject = (msg.subject or "").strip()
        sender  = (msg.sender  or "").strip()
        to      = (msg.to      or "").strip()
        date    = str(msg.date) if msg.date else ""
        body    = (msg.body or "").strip()

        # Strip URL noise + collapse whitespace
        body = _re.sub(r"https?://\S+", "", body)
        body = _re.sub(r"\s{3,}", "\n\n", body).strip()

        # Attachments list (names only — we don't recurse into them)
        atts = []
        try:
            atts = [a.longFilename for a in msg.attachments if a.longFilename]
        except Exception:
            pass

        header = []
        if subject: header.append(f"Subject: {subject}")
        if sender:  header.append(f"From:    {sender}")
        if to:      header.append(f"To:      {to}")
        if date:    header.append(f"Date:    {date}")
        if atts:    header.append(f"Attachments: {', '.join(atts)}")

        parts = ["\n".join(header)] if header else []
        if body:
            parts.append("---\n" + body)

        return "\n\n".join(parts) or "[Empty email]"
    finally:
        try:
            os.unlink(tmp_path)
        except Exception:
            pass


# ── EML (plain email) ─────────────────────────────────────────────────────────

def _parse_eml(data: bytes) -> str:
    """Parse .eml using Python's stdlib email module."""
    import email
    from email import policy
    import re as _re

    msg = email.message_from_bytes(data, policy=policy.default)
    subject = msg.get("Subject", "").strip()
    sender  = msg.get("From",    "").strip()
    to      = msg.get("To",      "").strip()
    date    = msg.get("Date",    "").strip()

    body = ""
    if msg.is_multipart():
        for part in msg.walk():
            if part.get_content_type() == "text/plain":
                body = part.get_content()
                break
        if not body:
            for part in msg.walk():
                if part.get_content_type() == "text/html":
                    body = part.get_content()
                    break
    else:
        body = msg.get_content() if msg.get_content_type().startswith("text/") else ""

    body = _re.sub(r"https?://\S+", "", body or "")
    body = _re.sub(r"\s{3,}", "\n\n", body).strip()

    header = []
    if subject: header.append(f"Subject: {subject}")
    if sender:  header.append(f"From:    {sender}")
    if to:      header.append(f"To:      {to}")
    if date:    header.append(f"Date:    {date}")

    parts = ["\n".join(header)] if header else []
    if body:
        parts.append("---\n" + body)

    return "\n\n".join(parts) or "[Empty email]"


# ── CSV ───────────────────────────────────────────────────────────────────────

def _parse_csv(data: bytes) -> str:
    import csv
    text   = data.decode("utf-8", errors="replace")
    reader = csv.reader(io.StringIO(text))
    lines  = []
    for i, row in enumerate(reader):
        if i >= 100:
            break
        lines.append(" | ".join(row))
    return "\n".join(lines)


# ── Images (Claude vision) ────────────────────────────────────────────────────

def _parse_image(data: bytes, mime_type: str, api_key: str) -> str:
    try:
        import anthropic
        client  = anthropic.Anthropic(api_key=api_key)
        b64     = base64.standard_b64encode(data).decode("utf-8")
        message = client.messages.create(
            model="claude-haiku-4-5",
            max_tokens=1000,
            messages=[{
                "role": "user",
                "content": [
                    {
                        "type": "image",
                        "source": {"type": "base64", "media_type": mime_type, "data": b64},
                    },
                    {
                        "type": "text",
                        "text": (
                            "This is an attachment from a customer request for an electrical "
                            "control cabinet. Extract all technical information visible: "
                            "power ratings, motor counts, IP ratings, dimensions, component "
                            "references, wiring details, or any specs. Be concise and structured."
                        ),
                    },
                ],
            }],
        )
        return message.content[0].text
    except Exception as e:
        return f"[Image could not be parsed: {e}]"


# ── Plain text ────────────────────────────────────────────────────────────────

def _parse_txt(data: bytes) -> str:
    return data.decode("utf-8", errors="replace")


# ── Public API ────────────────────────────────────────────────────────────────

IMAGE_MIMES = {
    ".png":  "image/png",
    ".jpg":  "image/jpeg",
    ".jpeg": "image/jpeg",
    ".webp": "image/webp",
    ".gif":  "image/gif",
}


def parse_file(filename: str, data: bytes, api_key: str = "") -> dict:
    """
    Parse a single uploaded file and return:
      {name, type, text, error, char_count}
    """
    ext  = Path(filename).suffix.lower()
    name = Path(filename).name

    if ext not in SUPPORTED_EXTENSIONS:
        return {"name": name, "type": ext, "text": "", "error": f"Unsupported file type: {ext}", "char_count": 0}

    try:
        if ext == ".pdf":
            raw = _parse_pdf(data)
        elif ext in (".docx", ".doc"):
            raw = _parse_docx(data)
        elif ext in (".xlsx", ".xlsm", ".xls"):
            raw = _parse_xlsx(data)
        elif ext == ".msg":
            raw = _parse_msg(data)
        elif ext == ".eml":
            raw = _parse_eml(data)
        elif ext == ".csv":
            raw = _parse_csv(data)
        elif ext in IMAGE_MIMES:
            mime = IMAGE_MIMES[ext]
            raw  = _parse_image(data, mime, api_key)
        else:
            raw = _parse_txt(data)

        raw = raw.strip()
        clipped = raw[:MAX_TEXT_CHARS]
        if len(raw) > MAX_TEXT_CHARS:
            clipped += f"\n... [truncated — {len(raw)} chars total]"

        return {"name": name, "type": ext.lstrip("."), "text": clipped, "error": None, "char_count": len(clipped)}

    except Exception as e:
        return {"name": name, "type": ext.lstrip("."), "text": "", "error": str(e), "char_count": 0}


def build_attachments_context(parsed_files: list) -> str:
    """
    Build a context block to inject into the LLM prompt.
    Only includes files that parsed successfully.
    """
    sections = []
    for f in parsed_files:
        if f.get("text") and not f.get("error"):
            sections.append(
                f"=== Attached file: {f['name']} ===\n{f['text']}"
            )
    if not sections:
        return ""
    return "\n\n".join(sections)
